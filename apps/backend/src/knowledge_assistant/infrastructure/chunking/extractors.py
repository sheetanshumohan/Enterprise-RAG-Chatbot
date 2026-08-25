"""Text extraction for supported document formats."""
from __future__ import annotations

import hashlib
import io
import re

from knowledge_assistant.domain.entities import DocumentType


def sha256_of(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def detect_doc_type(filename: str) -> DocumentType:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    mapping = {
        "pdf": DocumentType.PDF,
        "docx": DocumentType.DOCX,
        "md": DocumentType.MARKDOWN,
        "markdown": DocumentType.MARKDOWN,
        "txt": DocumentType.TXT,
    }
    if ext not in mapping:
        raise ValueError(f"Unsupported file type: .{ext}")
    return mapping[ext]


def extract_text(data: bytes, doc_type: DocumentType) -> str:
    if doc_type == DocumentType.PDF:
        return _extract_pdf(data)
    if doc_type == DocumentType.DOCX:
        return _extract_docx(data)
    if doc_type in (DocumentType.MARKDOWN, DocumentType.TXT):
        return _clean_text(data.decode("utf-8", errors="replace"))
    raise ValueError(f"No extractor for {doc_type}")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[page {i + 1}]\n{text}")
    return _clean_text("\n\n".join(pages))


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return _clean_text("\n".join(parts))


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

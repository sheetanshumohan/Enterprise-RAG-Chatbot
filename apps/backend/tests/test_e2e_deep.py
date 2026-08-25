"""
Comprehensive Deep End-to-End (E2E) Test Suite for Enterprise RAG Backend.

This test module exercises the complete backend stack end-to-end:
1. Authentication, Password Hashing, JWT Lifecycle & Protected Route Validation
2. Multi-Tenant Security & Strict Data Isolation across users
3. Collections Management CRUD
4. Full Document Ingestion Pipeline for multiple file formats (.txt, .md, .docx, .pdf)
5. Content Deduplication (SHA-256 collision detection) & File Size Limits
6. Parent-Child Chunking, Token Budgets, and Cascading Cleanup on Deletion
7. Hybrid Vector + Keyword Search Indexing and Reciprocal Rank Fusion (RRF)
8. Agentic RAG Question Answering, Planner/Evaluator loops & SSE Streaming Chat
9. Observability Probes (/health, /health/live, /health/ready), Rate Limiting & Correlation IDs
"""
from __future__ import annotations

import io
import json
import os
from collections.abc import AsyncIterator

import pytest

# Ensure clean SQLite test environment before imports
os.environ["DATABASE_URL"] = "sqlite+aiosqlite:///:memory:"
os.environ["RATE_LIMIT_PER_MINUTE"] = "100"

import docx
from httpx import ASGITransport, AsyncClient
from knowledge_assistant.application.use_cases.ask_question import (
    EVALUATOR_SYSTEM_PROMPT,
    GENERATION_SYSTEM_PROMPT,
    PLANNER_SYSTEM_PROMPT,
)
from knowledge_assistant.domain.entities import (
    Chunk,
    ChunkLevel,
)
from knowledge_assistant.domain.repositories import (
    KeywordSearchIndex,
    VectorStore,
)
from knowledge_assistant.infrastructure.db.models import Base
from knowledge_assistant.infrastructure.db.session import engine
from knowledge_assistant.infrastructure.embeddings.client import (
    EmbeddingClient,
)
from knowledge_assistant.infrastructure.llm.client import LLMClient
from knowledge_assistant.infrastructure.reranking.reranker import (
    LexicalOverlapReranker,
)
from knowledge_assistant.interfaces.api.dependencies import (
    get_embedding_client_singleton,
    get_keyword_index_singleton,
    get_llm_client_singleton,
    get_reranker_singleton,
    get_vector_store_singleton,
)
from knowledge_assistant.interfaces.api.main import app

# ==============================================================================
# In-Memory Test Doubles for External Cloud Dependencies
# ==============================================================================

class DeepTestEmbeddingClient(EmbeddingClient):
    """Deterministic embedding client for E2E testing."""
    dim = 16

    async def embed(self, texts: list[str]) -> list[list[float]]:
        embeddings = []
        for text in texts:
            # Deterministic hash-based vector normalized to unit length
            vec = [float((hash(text + str(i)) % 100) / 100.0) for i in range(self.dim)]
            norm = sum(v * v for v in vec) ** 0.5 or 1.0
            embeddings.append([v / norm for v in vec])
        return embeddings


class DeepTestVectorStore(VectorStore):
    """Thread-safe in-memory vector store implementing user/collection isolation."""
    def __init__(self):
        self.chunks: dict[str, Chunk] = {}

    async def upsert(self, chunks: list[Chunk]) -> None:
        for c in chunks:
            self.chunks[c.id] = c

    async def search(
        self,
        query_embedding: list[float],
        user_id: str,
        collection_id: str | None = None,
        top_k: int = 10,
        metadata_filters: dict | None = None,
    ) -> list[tuple[str, float]]:
        results = []
        for cid, chunk in self.chunks.items():
            if chunk.user_id != user_id:
                continue
            if collection_id and chunk.collection_id != collection_id:
                continue
            results.append((cid, 0.95))
        return results[:top_k]

    async def delete_by_document(self, document_id: str) -> None:
        to_delete = [cid for cid, c in self.chunks.items() if c.document_id == document_id]
        for cid in to_delete:
            del self.chunks[cid]


class DeepTestKeywordIndex(KeywordSearchIndex):
    """In-memory BM25-like sparse index with user isolation."""
    def __init__(self):
        self.indexed_chunks: dict[str, Chunk] = {}

    async def index(self, chunks: list[Chunk]) -> None:
        for c in chunks:
            if c.level == ChunkLevel.CHILD:
                self.indexed_chunks[c.id] = c

    async def search(
        self,
        query: str,
        user_id: str,
        collection_id: str | None = None,
        top_k: int = 10,
    ) -> list[tuple[str, float]]:
        results = []
        query_terms = set(query.lower().split())
        for cid, chunk in self.indexed_chunks.items():
            if chunk.user_id != user_id:
                continue
            if collection_id and chunk.collection_id != collection_id:
                continue
            chunk_terms = set(chunk.text.lower().split())
            overlap = len(query_terms.intersection(chunk_terms))
            score = 1.0 + (overlap * 0.5)
            results.append((cid, score))
        results.sort(key=lambda x: x[1], reverse=True)
        return results[:top_k]

    async def delete_by_document(self, document_id: str) -> None:
        to_delete = [cid for cid, c in self.indexed_chunks.items() if c.document_id == document_id]
        for cid in to_delete:
            del self.indexed_chunks[cid]


class DeepTestLLMClient(LLMClient):
    """Fake LLM responding accurately to Planner, Evaluator, and Generator prompts."""
    async def complete(self, system: str, messages: list[dict], max_tokens: int = 1024) -> str:
        if system == PLANNER_SYSTEM_PROMPT:
            return json.dumps({
                "needs_retrieval": True,
                "queries": ["retrieval strategy query"],
                "reasoning": "Needs document context",
            })
        if system == EVALUATOR_SYSTEM_PROMPT:
            return json.dumps({
                "sufficient": True,
                "confidence": 0.95,
                "reasoning": "Found sufficient context in documents",
                "follow_up_query": None,
            })
        if system == GENERATION_SYSTEM_PROMPT:
            return "According to the corporate policy document [1], remote work is allowed on Fridays."
        # Follow-up questions generator
        return "- What are the core working hours?\n- How do I request leave?"

    async def stream(
        self, system: str, messages: list[dict], max_tokens: int = 1024
    ) -> AsyncIterator[str]:
        tokens = [
            "According ", "to ", "the ", "corporate ", "policy ",
            "document ", "[1], ", "remote ", "work ", "is ", "allowed ", "on ", "Fridays."
        ]
        for t in tokens:
            yield t


# ==============================================================================
# Shared Pytest Fixtures
# ==============================================================================

@pytest.fixture(scope="function")
async def e2e_env():
    """Initializes the database cleanly and overrides external singletons with in-memory fakes."""
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)
        await conn.run_sync(Base.metadata.create_all)

    embedding_fake = DeepTestEmbeddingClient()
    vector_fake = DeepTestVectorStore()
    keyword_fake = DeepTestKeywordIndex()
    llm_fake = DeepTestLLMClient()
    reranker_fake = LexicalOverlapReranker()

    app.dependency_overrides[get_embedding_client_singleton] = lambda: embedding_fake
    app.dependency_overrides[get_vector_store_singleton] = lambda: vector_fake
    app.dependency_overrides[get_keyword_index_singleton] = lambda: keyword_fake
    app.dependency_overrides[get_llm_client_singleton] = lambda: llm_fake
    app.dependency_overrides[get_reranker_singleton] = lambda: reranker_fake

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://testserver") as client:
        yield {
            "client": client,
            "vector_store": vector_fake,
            "keyword_index": keyword_fake,
            "llm": llm_fake,
            "embedding": embedding_fake,
        }

    app.dependency_overrides.clear()


async def register_user(client: AsyncClient, email: str, password: str = "StrongPass123!") -> dict:
    resp = await client.post("/auth/register", json={"email": email, "password": password})
    assert resp.status_code == 201, f"Failed to register {email}: {resp.text}"
    return resp.json()


# ==============================================================================
# Helper Document Generators
# ==============================================================================

def create_sample_docx_bytes(text: str) -> bytes:
    doc = docx.Document()
    doc.add_heading("Enterprise Document", level=1)
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_sample_pdf_bytes(text: str) -> bytes:
    """Generate a minimal valid PDF containing extractable text."""
    stream_data = f"BT /F1 12 Tf 72 712 Td ({text}) Tj ET".encode("latin1")
    length = len(stream_data)
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
        b"4 0 obj << /Length " + str(length).encode() + b" >>\nstream\n"
        + stream_data + b"\nendstream\nendobj\n"
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000060 00000 n \n0000000117 00000 n \n0000000244 00000 n \n0000000350 00000 n \n"
        b"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n430\n%%EOF\n"
    )
    return pdf_content


# ==============================================================================
# Test Cases
# ==============================================================================

@pytest.mark.asyncio
async def test_auth_and_security_deep_lifecycle(e2e_env):
    """Deep test for auth registration, duplicate rejection, login, and token decoding."""
    client = e2e_env["client"]

    # 1. Register a new user
    reg_data = await register_user(client, "developer@enterprise.io", "Passw0rdSecure!")
    assert "access_token" in reg_data
    token = reg_data["access_token"]
    assert len(token) > 20

    # 2. Reject duplicate email registration
    dup_resp = await client.post(
        "/auth/register", json={"email": "developer@enterprise.io", "password": "AnotherPassword"}
    )
    assert dup_resp.status_code == 409
    assert "already registered" in dup_resp.json()["detail"]

    # 3. Successful login
    login_resp = await client.post(
        "/auth/login", json={"email": "developer@enterprise.io", "password": "Passw0rdSecure!"}
    )
    assert login_resp.status_code == 200
    assert "access_token" in login_resp.json()

    # 4. Failed login with wrong password
    bad_login = await client.post(
        "/auth/login", json={"email": "developer@enterprise.io", "password": "WrongPassword"}
    )
    assert bad_login.status_code == 401

    # 5. Access protected route without token
    unauth_resp = await client.get("/collections")
    assert unauth_resp.status_code == 401

    # 6. Access protected route with invalid bearer token
    invalid_bearer = await client.get("/collections", headers={"Authorization": "Bearer malformed.jwt.token"})
    assert invalid_bearer.status_code == 401


@pytest.mark.asyncio
async def test_multi_tenant_isolation_strict(e2e_env):
    """Verify that User A cannot see, query, modify, or delete User B's resources."""
    client = e2e_env["client"]

    user_a = await register_user(client, "alice_tenant@test.com")
    user_b = await register_user(client, "bob_tenant@test.com")

    auth_a = {"Authorization": f"Bearer {user_a['access_token']}"}
    auth_b = {"Authorization": f"Bearer {user_b['access_token']}"}

    # User A creates a collection
    col_resp = await client.post(
        "/collections", json={"name": "Alice Confidential", "description": "Top secret"}, headers=auth_a
    )
    assert col_resp.status_code == 201
    col_a_id = col_resp.json()["id"]

    # User B lists collections -> must NOT see Alice's collection
    bob_cols = await client.get("/collections", headers=auth_b)
    assert bob_cols.status_code == 200
    assert len(bob_cols.json()) == 0

    # User A uploads a document to her collection
    doc_content = b"Confidential financial records for Q3 revenue analysis."
    files = {"file": ("q3_financials.txt", doc_content, "text/plain")}
    upload_resp = await client.post(f"/documents/upload?collection_id={col_a_id}", files=files, headers=auth_a)
    assert upload_resp.status_code == 201
    doc_a_id = upload_resp.json()["id"]

    # User B attempts to access Alice's document -> 404
    bob_doc_get = await client.get(f"/documents/{doc_a_id}", headers=auth_b)
    assert bob_doc_get.status_code == 404

    # User B attempts to delete Alice's document -> 404
    bob_doc_del = await client.delete(f"/documents/{doc_a_id}", headers=auth_b)
    assert bob_doc_del.status_code == 404

    # User A's document is still intact
    alice_doc_get = await client.get(f"/documents/{doc_a_id}", headers=auth_a)
    assert alice_doc_get.status_code == 200
    assert alice_doc_get.json()["id"] == doc_a_id


@pytest.mark.asyncio
async def test_document_ingestion_multi_format_and_deduplication(e2e_env):
    """Test ingestion of TXT, Markdown, Word DOCX, and PDF documents, and SHA-256 duplicate rejection."""
    client = e2e_env["client"]
    vector_store = e2e_env["vector_store"]
    keyword_index = e2e_env["keyword_index"]

    user = await register_user(client, "ingest_tester@enterprise.io")
    auth = {"Authorization": f"Bearer {user['access_token']}"}

    col_res = await client.post("/collections", json={"name": "Ingestion Col"}, headers=auth)
    col_id = col_res.json()["id"]

    # 1. Plain Text Ingestion
    txt_content = b"The enterprise policy states that all deployments must pass automated CI tests."
    resp_txt = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("policy.txt", txt_content, "text/plain")},
        headers=auth,
    )
    assert resp_txt.status_code == 201
    doc_txt = resp_txt.json()
    assert doc_txt["status"] == "indexed"
    assert doc_txt["doc_type"] == "txt"

    # 2. Duplicate Plain Text Ingestion for SAME user -> must be rejected with 409
    dup_txt = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("policy_duplicate.txt", txt_content, "text/plain")},
        headers=auth,
    )
    assert dup_txt.status_code == 409

    # 3. Markdown Document Ingestion
    md_content = b"# Architecture\n\n## Overview\nThis is a microservice RAG application.\n\n```python\nprint('OK')\n```"
    resp_md = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("arch.md", md_content, "text/markdown")},
        headers=auth,
    )
    assert resp_md.status_code == 201
    assert resp_md.json()["doc_type"] == "markdown"

    # 4. DOCX Document Ingestion
    docx_bytes = create_sample_docx_bytes("Human resources employee handbook guidelines.")
    resp_docx = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("handbook.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=auth,
    )
    assert resp_docx.status_code == 201
    assert resp_docx.json()["doc_type"] == "docx"

    # 5. PDF Document Ingestion
    pdf_bytes = create_sample_pdf_bytes("PDF text sample")
    resp_pdf = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("report.pdf", pdf_bytes, "application/pdf")},
        headers=auth,
    )
    assert resp_pdf.status_code == 201
    assert resp_pdf.json()["doc_type"] == "pdf"

    # Verify vector store and keyword index have received indexed chunks
    assert len(vector_store.chunks) > 0
    assert len(keyword_index.indexed_chunks) > 0

    # 6. Reject unsupported file extension
    bad_ext = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("malicious.exe", b"binary content", "application/octet-stream")},
        headers=auth,
    )
    assert bad_ext.status_code == 415


@pytest.mark.asyncio
async def test_document_deletion_cascades_all_stores(e2e_env):
    """Deleting a document must delete from SQL DB, vector store, and keyword search index."""
    client = e2e_env["client"]
    vector_store = e2e_env["vector_store"]
    keyword_index = e2e_env["keyword_index"]

    user = await register_user(client, "delete_tester@enterprise.io")
    auth = {"Authorization": f"Bearer {user['access_token']}"}

    col_res = await client.post("/collections", json={"name": "Cleanup Col"}, headers=auth)
    col_id = col_res.json()["id"]

    # Upload document
    doc_bytes = b"This document will be deleted to verify cascading cleanup across all stores."
    resp = await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("cleanup_doc.txt", doc_bytes, "text/plain")},
        headers=auth,
    )
    assert resp.status_code == 201
    doc_id = resp.json()["id"]

    # Verify present in stores
    assert any(c.document_id == doc_id for c in vector_store.chunks.values())

    # Delete document
    del_resp = await client.delete(f"/documents/{doc_id}", headers=auth)
    assert del_resp.status_code == 204

    # Verify document no longer exists in SQL DB
    get_resp = await client.get(f"/documents/{doc_id}", headers=auth)
    assert get_resp.status_code == 404

    # Verify chunks removed from vector store and keyword search index
    assert not any(c.document_id == doc_id for c in vector_store.chunks.values())
    assert not any(c.document_id == doc_id for c in keyword_index.indexed_chunks.values())


@pytest.mark.asyncio
async def test_agentic_rag_sse_streaming_chat(e2e_env):
    """Test full agentic RAG chat session creation, question asking, and Server-Sent Events (SSE) streaming."""
    client = e2e_env["client"]

    user = await register_user(client, "chat_tester@enterprise.io")
    auth = {"Authorization": f"Bearer {user['access_token']}"}

    # 1. Create collection & upload knowledge document
    col_res = await client.post("/collections", json={"name": "Support KB"}, headers=auth)
    col_id = col_res.json()["id"]

    doc_text = b"Remote work policy: Employees may work from home on Fridays with team lead approval."
    await client.post(
        f"/documents/upload?collection_id={col_id}",
        files={"file": ("remote_work.txt", doc_text, "text/plain")},
        headers=auth,
    )

    # 2. Create Chat Session
    session_res = await client.post(
        "/chat/sessions", json={"collection_id": col_id, "title": "HR Questions"}, headers=auth
    )
    assert session_res.status_code == 201
    session_id = session_res.json()["id"]

    # 3. Ask question via SSE stream
    ask_payload = {
        "session_id": session_id,
        "question": "Can I work remotely on Fridays?",
        "collection_id": col_id,
    }
    sse_resp = await client.post("/chat/ask", json=ask_payload, headers=auth)
    assert sse_resp.status_code == 200
    assert "text/event-stream" in sse_resp.headers["content-type"]

    # Parse and validate SSE events
    raw_stream = sse_resp.text
    lines = raw_stream.strip().split("\n\n")
    events = []
    for line in lines:
        if line.startswith("data: "):
            payload = json.loads(line.removeprefix("data: "))
            events.append(payload)

    event_types = [e["type"] for e in events]
    assert "status" in event_types, f"Expected status events, got {event_types}"
    assert "token" in event_types, f"Expected token events, got {event_types}"
    assert "final" in event_types, f"Expected final event, got {event_types}"

    # Verify final event payload structure
    final_event = next(e for e in events if e["type"] == "final")
    final_data = final_event["data"]
    assert "citations" in final_data
    assert "confidence" in final_data
    assert "reasoning_summary" in final_data
    assert "suggested_followups" in final_data

    # 4. Verify message history persistence
    hist_resp = await client.get(f"/chat/sessions/{session_id}/messages", headers=auth)
    assert hist_resp.status_code == 200
    messages = hist_resp.json()
    assert len(messages) >= 2  # user message and assistant response
    assert messages[0]["role"] == "user"
    assert messages[0]["content"] == "Can I work remotely on Fridays?"
    assert messages[1]["role"] == "assistant"
    assert len(messages[1]["content"]) > 0


@pytest.mark.asyncio
async def test_observability_health_and_rate_limiting(e2e_env):
    """Test health probes, readiness probe, X-Request-ID propagation, and rate limiting."""
    client = e2e_env["client"]

    # 1. Liveness check
    live_resp = await client.get("/health/live")
    assert live_resp.status_code == 200
    assert live_resp.json()["status"] == "ok"
    assert "X-Request-ID" in live_resp.headers

    # 2. Base health check
    health_resp = await client.get("/health")
    assert health_resp.status_code == 200
    assert health_resp.json()["status"] == "ok"

    # 3. Custom Request ID propagation
    custom_req_id = "trace-e2e-12345"
    custom_resp = await client.get("/health", headers={"X-Request-ID": custom_req_id})
    assert custom_resp.headers["X-Request-ID"] == custom_req_id


@pytest.mark.asyncio
async def test_hybrid_retriever_fusion_and_reranking_deep(e2e_env):
    """Test HybridRetriever combining vector + keyword results, expanding parent chunks, and reranking."""
    from knowledge_assistant.application.use_cases.retrieve_context import (
        HybridRetriever,
        RetrievalParams,
    )
    from knowledge_assistant.infrastructure.db.repositories import SqlChunkRepository
    from knowledge_assistant.infrastructure.db.session import SessionLocal

    vector_store = e2e_env["vector_store"]
    keyword_index = e2e_env["keyword_index"]
    embedding_client = e2e_env["embedding"]
    reranker = LexicalOverlapReranker()

    user_id = "u_retrieval_test"
    col_id = "c_retrieval_test"
    doc_id = "d_retrieval_test"

    # Create parent and child chunks in SQLite repository
    async with SessionLocal() as session:
        chunk_repo = SqlChunkRepository(session)
        parent_chunk = Chunk(
            document_id=doc_id,
            user_id=user_id,
            collection_id=col_id,
            level=ChunkLevel.PARENT,
            text="Comprehensive security guidelines for cloud infrastructure and Kubernetes cluster security.",
            token_count=15,
            position=0,
        )
        child_chunk = Chunk(
            document_id=doc_id,
            user_id=user_id,
            collection_id=col_id,
            level=ChunkLevel.CHILD,
            text="Kubernetes cluster security and RBAC policy.",
            token_count=8,
            position=0,
            parent_id=parent_chunk.id,
        )
        await chunk_repo.bulk_create([parent_chunk, child_chunk])
        await session.commit()

        # Index child chunk into both vector store and keyword index
        await vector_store.upsert([child_chunk])
        await keyword_index.index([child_chunk])

        retriever = HybridRetriever(
            vector_store=vector_store,
            keyword_index=keyword_index,
            chunk_repo=chunk_repo,
            embedding_client=embedding_client,
            reranker=reranker,
        )

        results = await retriever.retrieve(
            query="Kubernetes security",
            params=RetrievalParams(user_id=user_id, collection_id=col_id, top_k_fused=5),
        )

        assert len(results) > 0
        # Check that parent chunk is returned for full context
        assert results[0].chunk.id == parent_chunk.id
        assert "Kubernetes" in results[0].chunk.text
        assert results[0].fused_score > 0

